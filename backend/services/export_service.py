"""
Export Service — Xuất đề thi thành file .docx chuẩn Bộ GD&ĐT
Kế thừa logic từ generate_word_template.py hiện có
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from typing import Dict, Any, List


def build_docx(exam: Dict[str, Any]) -> bytes:
    """
    Tạo file .docx theo chuẩn định dạng Bộ GD&ĐT.
    Font: Times New Roman 12pt, lề chuẩn.
    """
    config    = exam.get("config", {})
    questions = exam.get("questions", [])
    title     = exam.get("title", "Đề thi")

    mcq_questions   = [q for q in questions if q.get("type") == "Trắc nghiệm"]
    essay_questions = [q for q in questions if q.get("type") == "Tự luận"]

    doc = Document()

    # ── Lề trang ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # ── Header: Sở GD / Tên đề thi ────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    _remove_table_borders(header_table)

    cell_left = header_table.cell(0, 0)
    _add_run(cell_left.paragraphs[0], config.get("schoolName", "SỞ GIÁO DỤC VÀ ĐÀO TẠO"), bold=False, size=12)
    _add_run(cell_left.add_paragraph(), "ĐỀ THI THAM KHẢO", bold=True, size=12)
    for p in cell_left.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_right = header_table.cell(0, 1)
    _add_run(cell_right.paragraphs[0], f"KỲ THI: {title.upper()}", bold=True, size=12)
    _add_run(cell_right.add_paragraph(), f"Môn: {config.get('subject', '')}  —  Năm học: {config.get('schoolYear', '')}", size=12)
    _add_run(cell_right.add_paragraph(), f"Thời gian làm bài: {config.get('duration', 90)} phút", size=12)
    for p in cell_right.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # ── Phần I: Trắc nghiệm ───────────────────────────────────────────
    if mcq_questions:
        mcq_total_pts = sum(q.get("points", 0.25) for q in mcq_questions)
        p = doc.add_paragraph()
        _add_run(p, f"PHẦN I. TRẮC NGHIỆM KHÁCH QUAN ({mcq_total_pts:.1f} điểm)", bold=True, size=12)

        for i, q in enumerate(mcq_questions, 1):
            p_q = doc.add_paragraph()
            _add_run(p_q, f"Câu {i}: ", bold=True, size=12)
            _add_run(p_q, q.get("content", ""), size=12)

            if q.get("options"):
                opts = q["options"]
                # Hiển thị 4 phương án trên 1-2 dòng
                opt_text = "     ".join([f"{o['key']}. {o['text']}" for o in opts])
                p_opts = doc.add_paragraph()
                _add_run(p_opts, opt_text, size=12)

        doc.add_paragraph()

    # ── Phần II: Tự luận ──────────────────────────────────────────────
    if essay_questions:
        essay_total_pts = sum(q.get("points", 2.0) for q in essay_questions)
        p = doc.add_paragraph()
        _add_run(p, f"PHẦN II. TỰ LUẬN ({essay_total_pts:.1f} điểm)", bold=True, size=12)

        for i, q in enumerate(essay_questions, 1):
            pts = q.get("points", 2.0)
            p_q = doc.add_paragraph()
            _add_run(p_q, f"Câu {i} ({pts:.1f} điểm): ", bold=True, size=12)
            _add_run(p_q, q.get("content", ""), size=12)
            doc.add_paragraph()

    # ── Trang đáp án (page break) ─────────────────────────────────────
    doc.add_page_break()

    p_ans_title = doc.add_paragraph()
    p_ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_ans_title, "HƯỚNG DẪN CHẤM VÀ ĐÁP ÁN", bold=True, size=14)

    # Bảng đáp án trắc nghiệm
    if mcq_questions:
        p = doc.add_paragraph()
        _add_run(p, "I. TRẮC NGHIỆM", bold=True, size=12)

        col_count = len(mcq_questions) + 1
        ans_table = doc.add_table(rows=2, cols=col_count)
        ans_table.style = "Table Grid"

        hdr = ans_table.rows[0].cells
        ans = ans_table.rows[1].cells
        hdr[0].text = "Câu"
        ans[0].text = "Đáp án"

        for i, q in enumerate(mcq_questions):
            hdr[i + 1].text = str(i + 1)
            ans[i + 1].text = str(q.get("answer", ""))

        _format_table_font(ans_table)
        doc.add_paragraph()

    # Đáp án tự luận
    if essay_questions:
        p = doc.add_paragraph()
        _add_run(p, "II. TỰ LUẬN", bold=True, size=12)
        for i, q in enumerate(essay_questions, 1):
            p_q = doc.add_paragraph()
            _add_run(p_q, f"Câu {i}: ", bold=True, size=12)
            _add_run(p_q, q.get("explanation", q.get("answer", "")), size=12)

    # ── Serialize sang bytes ──────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helpers ───────────────────────────────────────────────────────────────

def _add_run(paragraph, text: str, bold: bool = False, size: int = 12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    # Đặt font cho cả East Asian (tiếng Việt)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    rPr.insert(0, rFonts)
    return run


def _remove_table_borders(table):
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.append(_OE("w:tblPr"))
    tblBorders = _OE("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = _OE(f"w:{edge}")
        el.set(_qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _format_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
