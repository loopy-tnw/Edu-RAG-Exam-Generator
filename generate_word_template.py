from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_math_exam_template(output_path):
    doc = Document()
    
    # 1. Header (Sở Giáo dục, Trường, Tên kỳ thi)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_left_1 = p_left.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO\nTP. CẦN THƠ\n")
    run_left_1.font.name = 'Times New Roman'
    run_left_1.font.size = Pt(12)
    run_left_2 = p_left.add_run("ĐỀ THI THAM KHẢO")
    run_left_2.font.name = 'Times New Roman'
    run_left_2.font.size = Pt(12)
    run_left_2.bold = True
    
    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_right_1 = p_right.add_run("KỲ THI TUYỂN SINH LỚP 10 THPT\nNĂM HỌC 2026 - 2027\n")
    run_right_1.font.name = 'Times New Roman'
    run_right_1.font.size = Pt(12)
    run_right_1.bold = True
    run_right_2 = p_right.add_run("Môn thi: TOÁN\nThời gian làm bài: 120 phút")
    run_right_2.font.name = 'Times New Roman'
    run_right_2.font.size = Pt(12)
    
    doc.add_paragraph() # Spacer
    
    # 2. Phần I: Trắc nghiệm khách quan
    p_part1 = doc.add_paragraph()
    run_part1 = p_part1.add_run("PHẦN I. TRẮC NGHIỆM KHÁCH QUAN (3.0 điểm)")
    run_part1.font.name = 'Times New Roman'
    run_part1.font.size = Pt(12)
    run_part1.bold = True
    
    # Câu 1
    p_q1 = doc.add_paragraph()
    run_q1 = p_q1.add_run("Câu 1: ")
    run_q1.font.name = 'Times New Roman'
    run_q1.font.size = Pt(12)
    run_q1.bold = True
    
    run_q1_content = p_q1.add_run("Biểu thức $\sqrt{x-2}$ xác định khi và chỉ khi:")
    run_q1_content.font.name = 'Times New Roman'
    run_q1_content.font.size = Pt(12)
    
    # Answers Q1 (A, B, C, D)
    p_ans1 = doc.add_paragraph()
    run_ans1 = p_ans1.add_run("A. $x > 2$               B. $x \geq 2$               C. $x < 2$               D. $x \leq 2$")
    run_ans1.font.name = 'Times New Roman'
    run_ans1.font.size = Pt(12)
    
    # Câu 2
    p_q2 = doc.add_paragraph()
    run_q2 = p_q2.add_run("Câu 2: ")
    run_q2.bold = True
    run_q2.font.name = 'Times New Roman'
    run_q2.font.size = Pt(12)
    
    run_q2_content = p_q2.add_run("Hàm số $y = (m-1)x + 3$ đồng biến trên $\mathbb{R}$ khi:")
    run_q2_content.font.name = 'Times New Roman'
    run_q2_content.font.size = Pt(12)
    
    p_ans2 = doc.add_paragraph()
    run_ans2 = p_ans2.add_run("A. $m > 1$               B. $m < 1$               C. $m \geq 1$               D. $m \neq 1$")
    run_ans2.font.name = 'Times New Roman'
    run_ans2.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 3. Phần II: Tự luận
    p_part2 = doc.add_paragraph()
    run_part2 = p_part2.add_run("PHẦN II. TỰ LUẬN (7.0 điểm)")
    run_part2.font.name = 'Times New Roman'
    run_part2.font.size = Pt(12)
    run_part2.bold = True
    
    # Câu 1 Tự luận
    p_tq1 = doc.add_paragraph()
    run_tq1 = p_tq1.add_run("Câu 1 (2.0 điểm): ")
    run_tq1.bold = True
    run_tq1.font.name = 'Times New Roman'
    run_tq1.font.size = Pt(12)
    
    run_tq1_content = p_tq1.add_run("Giải hệ phương trình sau:\n")
    run_tq1_content.font.name = 'Times New Roman'
    run_tq1_content.font.size = Pt(12)
    
    run_tq1_math = p_tq1.add_run("   (1) $2x + y = 5$\n   (2) $x - 3y = -1$")
    run_tq1_math.font.name = 'Times New Roman'
    run_tq1_math.font.size = Pt(12)
    
    # Page Break for Answers
    doc.add_page_break()
    
    # 4. Bảng đáp án
    p_ans_title = doc.add_paragraph()
    p_ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ans_title = p_ans_title.add_run("HƯỚNG DẪN CHẤM VÀ ĐÁP ÁN")
    run_ans_title.font.name = 'Times New Roman'
    run_ans_title.font.size = Pt(14)
    run_ans_title.bold = True
    
    # Bảng đáp án trắc nghiệm
    p_mcq_title = doc.add_paragraph()
    run_mcq_title = p_mcq_title.add_run("I. TRẮC NGHIỆM")
    run_mcq_title.bold = True
    run_mcq_title.font.name = 'Times New Roman'
    run_mcq_title.font.size = Pt(12)
    
    table_mcq = doc.add_table(rows=2, cols=5)
    table_mcq.style = 'Table Grid'
    
    hdr_cells = table_mcq.rows[0].cells
    hdr_cells[0].text = 'Câu'
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    hdr_cells[3].text = '3'
    hdr_cells[4].text = '...'
    
    ans_cells = table_mcq.rows[1].cells
    ans_cells[0].text = 'Đáp án'
    ans_cells[1].text = 'B'
    ans_cells[2].text = 'A'
    ans_cells[3].text = '...'
    ans_cells[4].text = '...'
    
    # Set font cho bảng
    for row in table_mcq.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(output_path)
    print(f"File created at {output_path}")

if __name__ == "__main__":
    import sys
    create_math_exam_template(sys.argv[1])
